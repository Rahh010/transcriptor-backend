from flask import Flask, request, send_file
from flask_cors import CORS
from youtube_transcript_api import YouTubeTranscriptApi
from docx import Document
from urllib.parse import urlparse, parse_qs
import tempfile

app = Flask(__name__)
CORS(app)

def extract_video_id(url):
    query = urlparse(url)
    if query.hostname == 'youtu.be':
        return query.path[1:]
    if query.hostname in ('www.youtube.com', 'youtube.com'):
        if query.path == '/watch':
            return parse_qs(query.query).get('v', [None])[0]
    return None

def get_transcript_docx(links):
    doc = Document()
    doc.add_heading("Combined YouTube Transcripts", 0)

    for i, link in enumerate(links, 1):
        video_id = extract_video_id(link)
        if not video_id:
            doc.add_heading(f"Video {i}: Invalid URL", level=2)
            continue
        try:
            transcript = YouTubeTranscriptApi.get_transcript(video_id)
            doc.add_heading(f"Video {i}: {link}", level=2)
            for entry in transcript:
                doc.add_paragraph(f"[{entry['start']:.2f}s] {entry['text']}")
        except Exception as e:
            doc.add_heading(f"Video {i}: {link} (Error)", level=2)
            doc.add_paragraph(str(e))

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

@app.route('/upload', methods=['POST'])
def upload():
    file = request.files.get('file')
    if not file:
        return "No file uploaded", 400

    links = file.read().decode('utf-8').splitlines()
    links = [link.strip() for link in links if link.strip()]

    docx_path = get_transcript_docx(links)
    return send_file(docx_path, as_attachment=True, download_name="transcripts.docx")

if __name__ == '__main__':
    app.run(debug=True, port=5000)
